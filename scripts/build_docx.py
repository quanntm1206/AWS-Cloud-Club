from __future__ import annotations

import argparse
import re
from pathlib import Path

import yaml
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
DOCX_VI = ROOT / "docs/docx-vi"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
TABLE_FILL = "E8EEF5"
WARNING_FILL = "FFF2CC"
WARNING_BORDER = "C65911"
REPO_URL = "https://github.com/quanntm1206/AWS-Cloud-Club"


def set_cell_shading(cell: object, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()  # type: ignore[attr-defined]
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell: object, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    properties = cell._tc.get_or_add_tcPr()  # type: ignore[attr-defined]
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = margins.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_table_geometry(table: object, widths: list[float]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT  # type: ignore[attr-defined]
    table.autofit = False  # type: ignore[attr-defined]
    properties = table._tbl.tblPr  # type: ignore[attr-defined]
    width = properties.first_child_found_in("w:tblW")
    width.set(qn("w:w"), "9360")
    width.set(qn("w:type"), "dxa")
    indent = properties.first_child_found_in("w:tblInd")
    if indent is None:
        indent = OxmlElement("w:tblInd")
        properties.append(indent)
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid  # type: ignore[attr-defined]
    for child in list(grid):
        grid.remove(child)
    for inches in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(round(inches * 1440)))
        grid.append(column)
    for row in table.rows:  # type: ignore[attr-defined]
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            cell_properties = cell._tc.get_or_add_tcPr()
            cell_width = cell_properties.first_child_found_in("w:tcW")
            cell_width.set(qn("w:w"), str(round(widths[index] * 1440)))
            cell_width.set(qn("w:type"), "dxa")


def add_table(
    document: Document,
    headers: list[str],
    rows: list[list[str]],
    widths: list[float],
    *,
    prevent_row_splits: bool = False,
) -> object:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header = table.rows[0]
    header._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for index, text in enumerate(headers):
        set_cell_shading(header.cells[index], TABLE_FILL)
        paragraph = header.cells[index].paragraphs[0]
        run = paragraph.add_run(text)
        run.bold = True
    for values in rows:
        row = table.add_row()
        if prevent_row_splits:
            row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))  # type: ignore[attr-defined]
        cells = row.cells
        for index, text in enumerate(values):
            cells[index].text = text
    set_table_geometry(table, widths)
    document.add_paragraph("")
    return table


def set_labeled_cell(cell: object, entries: list[tuple[str, str]]) -> None:
    cell.text = ""  # type: ignore[attr-defined]
    for index, (label, value) in enumerate(entries):
        paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()  # type: ignore[attr-defined]
        paragraph.add_run(f"{label}: ").bold = True
        paragraph.add_run(value)


def add_warning(document: Document, title: str, body: str) -> None:
    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    set_cell_shading(cell, WARNING_FILL)
    properties = cell._tc.get_or_add_tcPr()
    borders = properties.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        properties.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "12")
        border.set(qn("w:color"), WARNING_BORDER)
        borders.append(border)
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(f"CẢNH BÁO CHI PHÍ - {title}\n")
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(WARNING_BORDER)
    paragraph.add_run(body)
    set_table_geometry(table, [6.5])
    document.add_paragraph("")


def configure_styles(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = section.bottom_margin = Inches(1)
    section.left_margin = section.right_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in (
        ("Title", 24, DARK_BLUE, 0, 12),
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = document.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    if "Cost Warning" not in document.styles:
        warning = document.styles.add_style("Cost Warning", WD_STYLE_TYPE.PARAGRAPH)
        warning.font.name = "Calibri"
        warning.font.size = Pt(11)
        warning.font.color.rgb = RGBColor.from_string(WARNING_BORDER)
        warning.font.bold = True
    for section in document.sections:
        header = section.header.paragraphs[0]
        header.text = "AWS Cloud Club | Machine Learning Engineer Roadmap"
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header.runs[0].font.size = Pt(8)
        header.runs[0].font.color.rgb = RGBColor(100, 100, 100)
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.add_run("AWS Cloud Club  |  Tài liệu học tập - kiểm tra nguồn trước mỗi cohort")
        footer.runs[0].font.size = Pt(8)


def set_document_language(document: Document, language: str) -> None:
    settings = document.settings.element
    theme_languages = settings.findall(qn("w:themeFontLang"))
    if theme_languages:
        theme_language = theme_languages[0]
        for duplicate in theme_languages[1:]:
            settings.remove(duplicate)
    else:
        theme_language = OxmlElement("w:themeFontLang")
        settings.append(theme_language)
    theme_language.set(qn("w:val"), language)
    theme_language.set(qn("w:eastAsia"), language)

    styles = document.styles.element
    defaults = styles.find(qn("w:docDefaults"))
    if defaults is None:
        defaults = OxmlElement("w:docDefaults")
        styles.insert(0, defaults)
    run_defaults = defaults.find(qn("w:rPrDefault"))
    if run_defaults is None:
        run_defaults = OxmlElement("w:rPrDefault")
        defaults.append(run_defaults)
    run_properties = run_defaults.find(qn("w:rPr"))
    if run_properties is None:
        run_properties = OxmlElement("w:rPr")
        run_defaults.append(run_properties)
    run_language = run_properties.find(qn("w:lang"))
    if run_language is None:
        run_language = OxmlElement("w:lang")
        run_properties.append(run_language)
    run_language.set(qn("w:val"), language)
    run_language.set(qn("w:eastAsia"), language)
    run_language.attrib.pop(qn("w:bidi"), None)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(item)
        paragraph.paragraph_format.left_indent = Inches(0.375)
        paragraph.paragraph_format.first_line_indent = Inches(-0.188)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.25


def add_hyperlink(paragraph: object, label: str, url: str) -> None:
    relationship_id = paragraph.part.relate_to(  # type: ignore[attr-defined]
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.extend([color, underline])
    text = OxmlElement("w:t")
    text.text = label
    run.extend([properties, text])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)  # type: ignore[attr-defined]


def add_bookmark(paragraph: object, name: str, bookmark_id: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)  # type: ignore[attr-defined]
    paragraph._p.append(end)  # type: ignore[attr-defined]


def add_internal_link(paragraph: object, label: str, anchor: str) -> None:
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.extend([color, underline])
    text = OxmlElement("w:t")
    text.text = label
    run.extend([properties, text])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)  # type: ignore[attr-defined]


def add_section_heading(document: Document, text: str, level: int, anchor: str, bookmark_id: int) -> object:
    heading = document.add_heading(text, level=level)
    add_bookmark(heading, anchor, bookmark_id)
    return heading


def extract_markdown_section(markdown: str, heading: str) -> str:
    return markdown.split(f"## {heading}", 1)[1].split("##", 1)[0].strip()


def markdown_items(section: str) -> list[str]:
    items: list[str] = []
    for line in section.splitlines():
        stripped = line.strip()
        if stripped.startswith("- "):
            items.append(stripped[2:].replace("`", ""))
        elif re.match(r"^\d+\. ", stripped):
            items.append(stripped.split(". ", 1)[1].replace("`", ""))
        elif stripped and items:
            items[-1] = f"{items[-1]} {stripped.replace('`', '')}"
    return items


def add_markdown_concept_groups(document: Document, section: str) -> None:
    groups = re.split(r"(?m)^###\s+", section)
    for group in groups[1:]:
        lines = group.strip().splitlines()
        if not lines:
            continue
        document.add_heading(lines[0].strip(), level=4)
        current_label = ""
        current_value: list[str] = []

        def flush() -> None:
            nonlocal current_label, current_value
            if not current_label:
                return
            paragraph = document.add_paragraph()
            paragraph.add_run(f"{current_label}: ").bold = True
            paragraph.add_run(" ".join(current_value).replace("`", "").replace("**", "").strip())
            current_label = ""
            current_value = []

        for line in lines[1:]:
            stripped = line.strip()
            match = re.match(r"^\*\*(.+?):\*\*\s*(.*)$", stripped)
            if match:
                flush()
                current_label = match.group(1)
                current_value = [match.group(2)]
            elif stripped and current_label:
                current_value.append(stripped)
        flush()


def add_command(document: Document, command: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.style = document.styles["Normal"]
    paragraph.paragraph_format.left_indent = Inches(0.22)
    paragraph.paragraph_format.right_indent = Inches(0.1)
    set_paragraph_shading(paragraph, "F3F6F8")
    run = paragraph.add_run(command)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def set_paragraph_shading(paragraph: object, fill: str) -> None:
    properties = paragraph._p.get_or_add_pPr()  # type: ignore[attr-defined]
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def build(output: Path) -> None:
    curriculum = yaml.safe_load((DOCX_VI / "curriculum/curriculum.yml").read_text(encoding="utf-8"))
    assessments = yaml.safe_load((DOCX_VI / "curriculum/assessment.yml").read_text(encoding="utf-8"))
    sources = yaml.safe_load((ROOT / "docs/sources.yml").read_text(encoding="utf-8"))["sources"]
    document = Document()
    document.core_properties.title = "Machine Learning Engineer Roadmap - AWS Cloud Club"
    document.core_properties.subject = "Roadmap 24 tuần, local-first, free-compute-first và AWS cost-safe"
    document.core_properties.author = "AWS Cloud Club"
    document.core_properties.keywords = "Machine Learning, MLOps, AWS, Colab, Kaggle, roadmap"
    set_document_language(document, "vi-VN")
    configure_styles(document)
    bookmark_id = 1
    document.add_heading("Machine Learning Engineer Roadmap", 0)
    subtitle = document.add_paragraph("AWS CLOUD CLUB  |  24 TUẦN  |  8-10 GIỜ/TUẦN")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True
    document.add_paragraph("Từ nền tảng Machine Learning đến ML Engineering và AWS capstone tối ưu chi phí.")
    add_warning(
        document,
        "ĐỌC TRƯỚC KHI DÙNG AWS",
        "Core path không dùng GPU, NAT Gateway, EC2 hay SageMaker runtime. AWS Budgets chỉ cảnh báo, "
        "không phải hard cap; billing có thể trễ. Luôn cleanup và residual scan trong cùng phiên lab.",
    )
    add_table(
        document,
        ["Thông tin", "Giá trị"],
        [
            ["Đối tượng", "Đã biết lập trình, chưa học Machine Learning"],
            ["Hình thức", "Tự học cá nhân trong AWS Cloud Club"],
            ["Training", "Local CPU; chọn Colab Free hoặc Kaggle Free cho CV"],
            ["Capstone", "Tabular local-first; AWS deploy tùy plan/credit; CV transfer learning mở rộng"],
            ["Ngày kiểm nguồn volatile", "12/08/2026; kiểm lại trước mỗi cohort"],
        ],
        [1.875, 4.625],
    )
    document.add_page_break()
    document.add_heading("Mục lục", level=1)
    for label, anchor in (
        ("Bắt đầu trong 15 phút", "quick-start"),
        ("1. Cách dùng và chuẩn đầu ra", "how-to-use"),
        ("2. Roadmap 24 tuần", "roadmap"),
        ("3. Colab Free và Kaggle Free", "free-compute"),
        ("4. AWS cost-safe capstone", "aws-capstone"),
        ("5. Lab, mốc năng lực và rubric", "assessment"),
        ("6. Tổng kết năng lực", "capability-summary"),
        ("7. Danh mục lab", "lab-directory"),
        ("8. Glossary và nguồn", "glossary"),
    ):
        paragraph = document.add_paragraph()
        add_internal_link(paragraph, label, anchor)

    add_section_heading(document, "Bắt đầu trong 15 phút", 1, "quick-start", bookmark_id)
    bookmark_id += 1
    document.add_paragraph(
        "Nếu đây là lần đầu bạn mở tài liệu, đừng cố đọc hết. Hãy tải bộ khung, kiểm tra môi trường "
        "và chạy lab đầu tiên. Khi thấy báo cáo môi trường xuất hiện, bạn đã sẵn sàng bước vào tuần 1."
    )
    paragraph = document.add_paragraph("Repo mẫu: ")
    add_hyperlink(paragraph, "Mở repository AWS-Cloud-Club", REPO_URL)
    document.add_paragraph(f"Địa chỉ: {REPO_URL}")
    add_command(document, f"git clone {REPO_URL}.git")
    document.add_paragraph("Windows PowerShell:")
    add_command(document, "pwsh scripts/setup.ps1 -Profile core")
    add_command(document, "pwsh scripts/check.ps1 -Scope bootstrap")
    add_command(document, ".venv\\Scripts\\python.exe scripts/run_lab.py --lab 0")
    document.add_paragraph("macOS hoặc Linux:")
    add_command(document, "bash scripts/setup.sh --profile core")
    add_command(document, "bash scripts/check.sh --scope bootstrap")
    add_command(document, ".venv/bin/python scripts/run_lab.py --lab 0")
    add_bullets(
        document,
        [
            "Đọc roadmap/weeks/week-01.md để biết mục tiêu và nhịp học tuần đầu.",
            "Mở labs/lab-00-environment-and-reproducibility/README.md để làm đúng từng bước.",
            "Không cần fork, commit, push hay nộp bài. Mọi kết quả được lưu cục bộ để bạn tự đối chiếu.",
        ],
    )

    add_section_heading(document, "1. Cách dùng và chuẩn đầu ra", 1, "how-to-use", bookmark_id)
    bookmark_id += 1
    document.add_paragraph(
        "GitHub chỉ dùng để clone/download repo mẫu do chủ repo phát hành. Kết quả học tập và minh chứng "
        "được lưu cục bộ để người học tự đánh giá."
    )
    add_bullets(
        document,
        [
            "Mỗi tuần: khoảng 2 giờ đọc có mục tiêu, 2 giờ thực hành có hướng dẫn, 3-4 giờ lab, "
            "1 giờ tự kiểm tra và 1 giờ ghi nhật ký học tập.",
            "Làm phần cốt lõi trước phần mở rộng. Bạn không cần sweep tham số hay mua compute "
            "để hoàn thành roadmap.",
            "Mỗi thí nghiệm nên lưu câu hỏi, baseline, cách chia dữ liệu, seed, metric, runtime, "
            "artifact và điều còn hạn chế.",
            "Mốc năng lực dùng sản phẩm chạy được để tự đánh giá; accuracy đơn lẻ không đủ.",
        ],
    )
    document.add_heading("Chuẩn đầu ra sau 24 tuần", level=2)
    add_bullets(
        document,
        [
            "Xây pipeline tabular chống leakage và đánh giá theo constraint.",
            "Xây CV transfer-learning baseline với CPU fallback và failure analysis.",
            "Đóng gói, test, phục vụ inference và quản lý artifact.",
            "Triển khai Lambda private invoke, quan sát log, cleanup và cost audit trên AWS.",
            "Trình bày hai bộ tổng kết năng lực có model card và reproduction guide.",
        ],
    )
    add_warning(
        document,
        "Khi bị kẹt / Khi mắc kẹt",
        "Thu nhỏ dữ liệu, chạy lại baseline, đọc lỗi từ dòng đầu và quay về mini profile. Nếu vẫn chưa rõ, "
        "ghi lại điều bạn kỳ vọng, điều thực tế xảy ra và thử nghiệm nhỏ nhất tiếp theo. "
        "Đừng mua thêm compute để che một lỗi chưa hiểu.",
    )

    add_section_heading(document, "2. Roadmap 24 tuần", 1, "roadmap", bookmark_id)
    bookmark_id += 1
    document.add_heading("Bản đồ 6 chặng", level=2)
    phase_rows: list[list[str]] = []
    for phase, label, purpose in (
        ("foundation", "Nền tảng dữ liệu và toán", "Hiểu dữ liệu, vector và gradient qua ví dụ nhỏ."),
        ("classical-ml", "ML cổ điển", "Xây baseline, chống leakage và đánh giá đáng tin cậy."),
        ("applied-ml", "ML thực hành", "So sánh model, phân tích lỗi, hoàn thiện mini-project."),
        ("engineering", "ML engineering", "Đưa notebook thành package, test, API và CI."),
        ("deep-learning", "Deep Learning và CV", "Train tiết kiệm, lưu checkpoint, hiểu failure cases."),
        ("aws-capstone", "AWS capstone", "Tích hợp serverless ngắn hạn, dọn sạch, kiểm chi phí."),
    ):
        phase_weeks = [str(week["id"]) for week in curriculum["weeks"] if week["phase"] == phase]
        phase_rows.append([label, f"{phase_weeks[0]}-{phase_weeks[-1]}", purpose])
    add_table(document, ["Chặng", "Tuần", "Bạn sẽ đi được đến đâu"], phase_rows, [2.0, 0.65, 3.85])
    for week in curriculum["weeks"]:
        heading = document.add_heading(f"Tuần {week['id']:02d} - {week['title']}", level=2)
        # Week 1 follows the phase map; week 22 already flows to a fresh page after the dense week 21 spread.
        # Forcing either break leaves an almost empty or fully blank page in Word's paginator.
        heading.paragraph_format.page_break_before = week["id"] not in {1, 22}
        add_bookmark(heading, f"week-{week['id']:02d}", bookmark_id)
        bookmark_id += 1
        environment = ", ".join(week["environments"])
        add_table(
            document,
            ["Mục", "Chi tiết"],
            [
                ["Môi trường", environment],
                ["Workload", f"{week['hours']} giờ; core trước stretch"],
                ["Lab", week["lab"]],
                ["Mốc năng lực", week["milestone"] or "Không"],
                ["Cost class", week["cost_class"]],
            ],
            [1.181, 5.319],
        )
        week_doc = (DOCX_VI / f"roadmap/weeks/week-{week['id']:02d}.md").read_text(encoding="utf-8")
        goal = extract_markdown_section(week_doc, "Mục tiêu tuần")
        lead = document.add_paragraph()
        lead.add_run("Tuần này bạn sẽ ").bold = True
        lead.add_run(goal[:1].lower() + goal[1:])
        core = extract_markdown_section(week_doc, "Kiến thức cốt lõi")
        why = extract_markdown_section(week_doc, "Vì sao tuần này quan trọng")
        document.add_heading("Vì sao phần này quan trọng", level=3)
        document.add_paragraph(why.replace("**", ""))
        document.add_heading("Điều cần hiểu", level=3)
        add_bullets(document, markdown_items(core))
        vocabulary = extract_markdown_section(week_doc, "Từ khóa tuần này")
        document.add_heading("Từ khóa tuần này", level=3)
        for line in vocabulary.splitlines():
            stripped = line.strip()
            if not stripped:
                continue
            paragraph = document.add_paragraph()
            if stripped.startswith("**") and ":**" in stripped:
                label, value = stripped.split(":**", 1)
                paragraph.add_run(label.removeprefix("**") + ": ").bold = True
                paragraph.add_run(value.strip().replace("`", ""))
            else:
                paragraph.add_run(stripped.replace("`", "").replace("**", ""))
        concept_walkthrough = extract_markdown_section(week_doc, "Giải thích khái niệm")
        document.add_heading("Giải thích khái niệm", level=3)
        add_markdown_concept_groups(document, concept_walkthrough)
        prior_connections = extract_markdown_section(week_doc, "Kết nối kiến thức cũ")
        document.add_heading("Kết nối kiến thức cũ", level=3)
        connection_items = markdown_items(prior_connections)
        if connection_items:
            add_bullets(document, connection_items)
        else:
            document.add_paragraph(prior_connections.replace("`", "").replace("**", ""))
        guided = extract_markdown_section(week_doc, "Guided practice")
        document.add_heading("Thực hành có hướng dẫn", level=3)
        add_bullets(document, markdown_items(guided))
        lab_section = extract_markdown_section(week_doc, "Lab")
        paragraph = document.add_paragraph()
        paragraph.add_run("Thực hành: ").bold = True
        paragraph.add_run(lab_section.replace("**", ""))
        outcome = extract_markdown_section(week_doc, "Kết quả hướng tới")
        paragraph = document.add_paragraph()
        paragraph.add_run("Kết quả hướng tới: ").bold = True
        paragraph.add_run(outcome)
        self_check = extract_markdown_section(week_doc, "Tự kiểm tra")
        document.add_heading("Bạn đã sẵn sàng chuyển tuần khi", level=3)
        checks = markdown_items(self_check)
        add_bullets(document, [f"Bạn có thể trả lời bằng lời của mình: {item}" for item in checks])
        understood = extract_markdown_section(week_doc, "Dấu hiệu bạn đã hiểu")
        document.add_paragraph(f"Dấu hiệu bạn đã hiểu: {understood}")
        stuck = extract_markdown_section(week_doc, "Khi mắc kẹt")
        paragraph = document.add_paragraph()
        paragraph.add_run("Nếu bạn bị kẹt: ").bold = True
        paragraph.add_run(stuck)
        if week["id"] >= 21:
            add_warning(
                document,
                f"TUẦN {week['id']}",
                "Pre-check -> Estimate -> Dry-run -> Deploy -> Verify -> Cleanup -> Residual scan -> Cost audit. "
                "Không đánh dấu hoàn thành nếu còn resource.",
            )
    add_section_heading(document, "3. Colab Free và Kaggle Free", 1, "free-compute", bookmark_id)
    bookmark_id += 1
    document.add_paragraph(
        "Bạn chỉ cần chọn một nền tảng, không cần dùng cả hai. Tài nguyên miễn phí không được bảo đảm; "
        "accelerator, quota và runtime có thể thay đổi. Notebook tự phát hiện CUDA và chuyển sang "
        "cpu-mini khi GPU không có."
    )
    document.add_paragraph(
        "Bắt đầu từ notebooks/colab-training.ipynb hoặc notebooks/kaggle-training.ipynb. Nếu tên notebook "
        "trong repo thay đổi, mở notebooks/README.md để chọn bản dành cho transfer learning."
    )
    add_table(
        document,
        ["Bước", "Colab Free", "Kaggle Free"],
        [
            ["Mở", "Upload từ repo mẫu đã clone; Copy to Drive", "Import/Upload notebook"],
            ["Dữ liệu", "Download nhỏ hoặc Drive optional", "Add Dataset; không dùng private token trong cell"],
            ["Compute", "GPU chỉ khi dùng; cpu-mini fallback", "Accelerator nếu có; cpu-mini fallback"],
            ["Training", "Frozen backbone; 3-5 epoch max", "Frozen backbone; 3-5 epoch max"],
            ["Export", "Download artifacts.zip", "Save Version rồi download output"],
            ["Kết thúc", "Disconnect and delete runtime", "Tắt accelerator/session"],
        ],
        [1.1, 2.7, 2.7],
    )
    document.add_heading("Notebook contract", level=2)
    add_bullets(
        document,
        [
            "Environment check; dependency verification; configuration và seed.",
            "Data acquisition có mini fallback; validation trước model.",
            "Baseline; training; evaluation và error analysis.",
            "Artifact manifest, checkpoint export và release runtime.",
        ],
    )
    document.add_heading("Khi runtime bị ngắt hoặc hết GPU", level=2)
    add_bullets(
        document,
        [
            "Giảm về cpu-mini để kiểm tra luồng chạy; đây là đường hoàn thành hợp lệ, không phải phương án kém hơn.",
            "Resume từ checkpoint gần nhất thay vì train lại từ đầu; tải checkpoint về máy sau mỗi phiên quan trọng.",
            "Nếu quota chưa quay lại, tiếp tục error analysis và model card trên CPU. Không mua gói trả phí "
            "chỉ để theo kịp lịch.",
        ],
    )
    document.add_heading("Bạn có thể kết thúc phiên khi", level=2)
    add_bullets(
        document,
        [
            "Artifact và checkpoint đã tải về máy; config, seed, metric và runtime đã được ghi lại.",
            "Colab đã Disconnect and delete runtime hoặc Kaggle đã tắt accelerator/session.",
            "Không có token hay credential nằm trong cell, output hoặc artifact tải xuống.",
        ],
    )

    add_section_heading(document, "4. AWS cost-safe capstone", 1, "aws-capstone", bookmark_id)
    bookmark_id += 1
    add_warning(
        document,
        "HARD GUARDRAILS",
        "Core chỉ IAM, S3, Lambda, CloudWatch Logs và Budgets. HTTP API optional, tắt mặc định. "
        "Không bật public HTTP API trong đường học cốt lõi. Cấm GPU, EC2, NAT Gateway, SageMaker "
        "notebook/training/endpoint và resource chạy nền ngoài allowlist.",
    )
    document.add_paragraph(
        "Trước khi tạo tài nguyên, đọc aws/README.md và xác nhận đúng account, Region, plan cùng số dư credit. "
        "Chương trình không giả định bạn có USD 200. Một số ưu đãi mới có thể cấp USD 100 ban đầu và cho phép "
        "kiếm thêm tối đa USD 100 qua hoạt động; điều kiện có thể đổi, vì vậy màn hình Billing "
        "của chính tài khoản là nguồn quyết định."
    )
    document.add_paragraph(
        "AWS Organizations và Control Tower không thuộc đường lab cốt lõi. Đừng bật chỉ để học capstone này: "
        "chúng mở rộng phạm vi quản trị, không giúp model tốt hơn và có thể kéo theo cấu hình ngoài allowlist."
    )
    add_table(
        document,
        ["Control", "Giới hạn"],
        [
            ["Region", "us-east-1 mặc định"],
            ["Tags", "Project, Owner, Environment=learning, ExpiresAt"],
            ["Artifact", "Mục tiêu <50 MB; hard limit 200 MB"],
            ["Lambda", "512 MB; 15 giây; reserved concurrency 1"],
            ["Logs", "Retention một ngày"],
            ["S3", "Public block + encryption + lifecycle bảy ngày"],
            ["Cleanup", "Dry-run mặc định; exact project ID; CloudFormation delete; residual scan"],
        ],
        [1.875, 4.625],
    )
    document.add_heading("AWS lab - lệnh thực thi", level=2)
    document.add_paragraph(
        "1) Tạo actual + forecast Budget alerts trên Billing console. 2) Chạy cost check. "
        "3) Chạy preflight với AcknowledgeBudgetConfigured. 4) Deploy private Lambda bằng "
        "portable_model.json. 5) Invoke valid/invalid events và kiểm log. 6) Cleanup. "
        "7) Residual scan CloudFormation/S3/Lambda/Logs/IAM và audit Billing."
    )
    document.add_paragraph(
        "Mở aws/README.md để lấy tham số hiện hành. Chạy lệnh theo đúng thứ tự, cùng một ProjectId:"
    )
    add_command(
        document,
        "$project='student01'\n"
        "$region='us-east-1'\n"
        "$artifact='.artifacts/churn-model/portable_model.json'\n"
        "$expires=(Get-Date).AddDays(1).ToString('yyyy-MM-dd')",
    )
    add_command(document, "pwsh aws/scripts/cost-check.ps1 -ProjectId $project -Region $region")
    add_command(
        document,
        "pwsh aws/scripts/preflight.ps1 -ProjectId $project -Region $region "
        "-ArtifactPath $artifact -AcknowledgeBudgetConfigured",
    )
    add_command(
        document,
        "pwsh aws/scripts/deploy.ps1 -ProjectId $project -Owner $project -ExpiresAt $expires "
        "-ArtifactPath $artifact -Region $region -AcknowledgeBudgetConfigured",
    )
    add_command(document, "pwsh aws/scripts/cleanup.ps1 -ProjectId $project -Region $region")
    add_command(
        document,
        "pwsh aws/scripts/cleanup.ps1 -ProjectId $project -Region $region -Execute "
        "-ConfirmProjectId $project",
    )
    add_command(
        document,
        "pwsh aws/scripts/residual-scan.ps1 -ProjectId $project -Region $region -Json",
    )
    add_warning(
        document,
        "STOP CONDITIONS",
        "Dừng khi sai account/Region, chưa có Budget alerts, artifact quá 200 MB, template có forbidden "
        "resource hoặc cleanup target không khớp exact project ID.",
    )
    document.add_heading("Capstone A - Tabular classification", level=2)
    add_bullets(
        document,
        [
            "Train local CPU; export sklearn artifact và portable logistic JSON.",
            "Upload portable_model.json lên S3; Lambda chỉ dùng Python stdlib + boto3.",
            "Private invoke là đường thực hành duy nhất; public HTTP API chỉ được phân tích trên sơ đồ, "
            "không triển khai.",
            "Cleanup và residual scan là gate bắt buộc.",
        ],
    )
    document.add_heading("Capstone B - Image classification", level=2)
    add_bullets(
        document,
        [
            "Train trên một trong Colab/Kaggle; frozen backbone baseline; optional unfreeze block cuối; "
            "CPU mini path.",
            "Pretrained normalization phải lấy từ ResNet18 weights; FakeData/random weights "
            "chỉ là execution smoke.",
            "Notebook lưu last checkpoint sau mọi epoch để resume và best checkpoint khi validation loss tốt hơn "
            "để đánh giá. Checkpoint gồm model, optimizer, epoch, best metric, history, config, seed và class mapping. "
            "Checkpoint/resume là cơ chế kỹ thuật để tiếp tục training, không phải mốc năng lực.",
            "Macro/per-class metric, confusion matrix và tối đa 20 failure cases; nếu ít hơn, "
            "export toàn bộ và ghi limitation.",
            "AWS chỉ artifact checksum/upload optional và architecture ADR; không endpoint CV.",
        ],
    )
    add_section_heading(document, "5. Lab, mốc năng lực và rubric", 1, "assessment", bookmark_id)
    bookmark_id += 1
    add_table(
        document,
        ["Nhóm", "Số lượng", "Gate"],
        [
            ["Lab", "21 lab", "Mini path, config, test, failure evidence"],
            ["Mốc năng lực", "6", "70/100; cuối kỳ 75/100"],
            ["Capstone", "2", "Không leakage/secret; AWS cleanup nếu áp dụng"],
        ],
        [1.5, 1.0, 4.0],
    )
    document.add_heading("Cách tự chấm", level=2)
    document.add_paragraph(
        "Bạn tự chấm từng tiêu chí theo bốn mức bên dưới, rồi nhân trọng số tiêu chí với hệ số của mức. "
        "Tổng tối đa được giới hạn ở 100. Điểm số giúp nhìn ra phần cần ôn; gate bắt buộc vẫn có quyền dừng mốc."
    )
    scale_rows = []
    for level in ("emerging", "developing", "proficient", "exemplary"):
        value = assessments["scale"][level]
        scale_rows.append([value["label"], str(value["factor"]), value["description"]])
    add_table(document, ["Mức", "Hệ số", "Cách hiểu"], scale_rows, [1.25, 0.65, 4.6])
    document.add_heading("Gate bắt buộc", level=2)
    add_bullets(
        document,
        [
            "Không data leakage: mọi transform học từ dữ liệu phải fit trên train.",
            "Không secret: credential, token, account ID và dữ liệu nhạy cảm không nằm trong "
            "code hay artifact.",
            "Mini run tái lập: một lệnh chạy lại được với config, seed và tolerance đã ghi.",
            "Nếu dùng AWS: cleanup hoàn tất và residual scan không còn tài nguyên đã biết.",
        ],
    )
    for index, milestone in enumerate(assessments["milestones"], start=1):
        document.add_heading(f"Mốc năng lực {index:02d} - tuần {milestone['week']}", level=2)
        document.add_paragraph(
            f"Trọng tâm: {milestone['focus']}. Điểm đạt: {milestone['pass_score']}/100. "
            "Minh chứng đạt mốc: README, code/notebook, test evidence, metric, learning log và limitations; "
            "lưu cục bộ để tự kiểm tra."
        )
        add_table(
            document,
            ["Tiêu chí", "Trọng số"],
            [[criterion, f"{weight}%"] for criterion, weight in milestone["criteria"].items()],
            [5.25, 1.25],
        )

    add_section_heading(document, "6. Tổng kết năng lực", 1, "capability-summary", bookmark_id)
    bookmark_id += 1
    add_bullets(
        document,
        [
            "README có problem, architecture, quickstart, metric, demo, limitations và cost note.",
            "Repository không chứa secret, raw data/model lớn hoặc notebook output dư thừa.",
            "Experiment report ghi baseline, controlled candidates, negative results và reproduction command.",
            "Model card ghi intended use, out-of-scope, subgroup/failure behavior, safety/privacy "
            "và artifact checksum.",
            "AWS evidence gồm preflight, cost manifest, deployment manifest, cleanup và zero-residual report.",
        ],
    )
    document.add_heading("Ba hướng đi tiếp trong 90 ngày", level=2)
    document.add_paragraph(
        "Model Engineering phù hợp nếu bạn thích dữ liệu và thí nghiệm; ML Platform/MLOps phù hợp nếu bạn "
        "thích package, test, CI và monitoring; Applied Computer Vision phù hợp nếu bạn muốn đào sâu transfer "
        "learning và phân tích lỗi ảnh. Xem roadmap/sau-24-tuan.md để dùng lịch 30-60-90 ngày."
    )
    add_section_heading(document, "7. Danh mục lab", 1, "lab-directory", bookmark_id)
    bookmark_id += 1
    document.add_paragraph(
        "Mỗi tuần, mở README của lab trước khi chạy. Lệnh chung dùng số lab trong bảng; output chỉ là "
        "bằng chứng tự kiểm tra, không phải bài nộp."
    )
    lab_rows: list[list[str]] = []
    for week in curriculum["weeks"]:
        lab_id = week["lab"]
        lab_dirs = sorted((ROOT / "labs").glob(f"{lab_id}-*"))
        path = (
            lab_dirs[0].relative_to(ROOT).as_posix() + "/README.md"
            if lab_dirs
            else f"labs/{lab_id}/README.md"
        )
        lab_number = int(lab_id.split("-")[1])
        lab_rows.append([f"{week['id']:02d}", lab_id, path, f"scripts/run_lab.py --lab {lab_number}"])
    add_table(document, ["Tuần", "Lab", "Đọc trước", "Lệnh"], lab_rows, [0.5, 0.8, 3.35, 1.85])

    glossary_heading = add_section_heading(document, "8. Glossary và nguồn", 1, "glossary", bookmark_id)
    glossary_heading.paragraph_format.page_break_before = True
    glossary = yaml.safe_load((DOCX_VI / "curriculum/glossary.yml").read_text(encoding="utf-8"))["terms"]
    document.add_paragraph(
        "Đừng học thuộc bảng này một lượt. Mỗi thuật ngữ được giới thiệu trong lab, dùng lại ở các lab sau và "
        "gắn với một evidence cụ thể. Ba khái niệm dễ nhầm: data validation kiểm dữ liệu, validation set là "
        "một phần dữ liệu, model validation là quá trình đánh giá/chọn quyết định."
    )
    glossary_table = add_table(
        document,
        [
            "Thuật ngữ",
            "Hiểu đơn giản và vì sao quan trọng",
            "Ví dụ, điểm dễ nhầm và tự kiểm tra",
            "Giới thiệu ở",
        ],
        [
            [
                item["term"],
                f"{item['meaning']} Vì sao quan trọng: {item['why_it_matters']}",
                (
                    f"Ví dụ: {item['example']} Dễ nhầm với: {item['common_confusion']} "
                    f"Tự kiểm tra: {item['self_check']}"
                ),
                f"Lab {int(item['introduced_in']):02d}",
            ]
            for item in glossary
        ],
        [1.15, 2.55, 2.3, 0.5],
        prevent_row_splits=True,
    )
    for row, item in zip(glossary_table.rows[1:], glossary, strict=True):
        set_labeled_cell(
            row.cells[1],
            [
                ("Hiểu đơn giản", item["meaning"]),
                ("Vì sao quan trọng", item["why_it_matters"]),
            ],
        )
        set_labeled_cell(
            row.cells[2],
            [
                ("Ví dụ", item["example"]),
                ("Dễ nhầm với", item["common_confusion"]),
                ("Tự kiểm tra", item["self_check"]),
            ],
        )
    document.add_heading("Nguồn chính thức", level=2)
    for source in sources:
        paragraph = document.add_paragraph()
        paragraph.add_run(f"{source['title']} ({source['verified_on']}): ").bold = True
        add_hyperlink(paragraph, f"Mở tài liệu {source['title']}", source["url"])
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--output", type=Path, default=ROOT / "dist/ML-Roadmap-24-Tuan-AWS-Cloud-Club.docx")
    args = parser.parse_args()
    build(args.output)
    print(f"DOCX BUILT: {args.output}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
