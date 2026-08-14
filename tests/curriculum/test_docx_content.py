import re
from pathlib import Path
from xml.etree import ElementTree
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn

from scripts.build_docx import build, markdown_items


def _extract_text(document: Document) -> str:
    return "\n".join(node.text or "" for node in document.element.iter(qn("w:t")))


def test_markdown_items_preserves_wrapped_list_content() -> None:
    section = """- First clause
  continues here and keeps `inline code`.
- Second item."""
    assert markdown_items(section) == [
        "First clause continues here and keeps inline code.",
        "Second item.",
    ]


def test_docx_contains_all_weeks_and_safety_language(tmp_path: Path) -> None:
    output = tmp_path / "roadmap.docx"
    build(output)
    document = Document(output)
    text = _extract_text(document)
    for week in range(1, 25):
        assert f"Tuần {week:02d}" in text
    assert "Budget không phải hard cap" in text or "Budgets chỉ cảnh báo" in text
    assert "Residual scan" in text or "residual scan" in text
    assert "Capstone A" in text and "Capstone B" in text
    assert "Mở tài liệu AWS Free Tier" in text
    assert "Thực hành:" in text and "Kết quả hướng tới:" in text
    assert "Mốc năng lực" in text
    assert text.count("Minh chứng đạt mốc") == 6
    assert "Tổng kết năng lực" in text
    assert "GitHub chỉ dùng để clone/download repo mẫu do chủ repo phát hành" in text
    assert text.count("GitHub") == 1
    assert "Nộp GitHub" not in text
    assert "Portfolio GitHub" not in text
    assert "Lab, checkpoint và rubric" not in text
    assert "checkpoint-01" not in text
    assert "Checkpoint/resume" in text
    assert "Checkpoint gồm model, optimizer" in text
    assert "AcknowledgeBudgetConfigured" in text
    assert "CloudFormation/S3/Lambda/Logs/IAM" in text
    assert "Không tạo API Gateway" in text
    assert "https://github.com/quanntm1206/AWS-Cloud-Club" in text
    assert "git clone https://github.com/quanntm1206/AWS-Cloud-Club.git" in text
    assert "scripts/setup.ps1" in text and "scripts/check.ps1" in text
    assert "scripts/run_lab.py --lab 0" in text
    assert "21 lab" in text
    assert "Khi mắc kẹt" in text
    assert "Dấu hiệu bạn đã hiểu" in text
    assert "AWS Organizations" in text and "Control Tower" in text
    assert "USD 100" in text and "tối đa USD 100" in text
    assert "Không bật public HTTP API" in text
    assert "Ba hướng đi tiếp trong 90 ngày" in text
    assert "Model Engineering" in text and "ML Platform/MLOps" in text
    assert "random-weight fallback chỉ smoke code và chưa đạt gate transfer learning" in text
    assert "USD 200 được cấp hết ngay. Free Plan kết thúc sau 6 tháng" in text
    assert "Free Plan tự nâng Paid Plan. Không dùng hai tính năng này" in text
    assert "nhưng không phải hard spending cap" in text
    assert "Tabular local-first; AWS deploy tùy plan/credit" in text
    assert "Tabular AWS bắt buộc" not in text
    assert "Từ khóa tuần này" in text
    assert "Thuật ngữ mới" in text and "Ôn lại" in text
    assert text.count("Giải thích khái niệm") == 24
    assert text.count("Kết nối kiến thức cũ") == 24
    assert "Vì sao quan trọng" in text
    assert "Dễ nhầm với" in text
    assert "dataset" in text and "augmentation" in text
    assert "Ví dụ" in text and "Giới thiệu ở" in text
    for command in (
        "pwsh aws/scripts/cost-check.ps1 -ProjectId $project -Region $region",
        "pwsh aws/scripts/preflight.ps1 -ProjectId $project -Region $region",
        "pwsh aws/scripts/deploy.ps1 -ProjectId $project -Owner $project",
        "pwsh aws/scripts/cleanup.ps1 -ProjectId $project -Region $region -Execute -ConfirmProjectId $project",
        "pwsh aws/scripts/residual-scan.ps1 -ProjectId $project -Region $region -Json",
    ):
        assert command in text
    assert "deploy.ps1 -WhatIf" not in text and "deploy.ps1 -Execute" not in text
    assert document.core_properties.title == "Machine Learning Engineer Roadmap - AWS Cloud Club"
    assert document.core_properties.author == "AWS Cloud Club"
    assert "local-first" in document.core_properties.subject
    with ZipFile(output) as archive:
        settings = archive.read("word/settings.xml").decode("utf-8")
        styles = archive.read("word/styles.xml")
        relationships = archive.read("word/_rels/document.xml.rels").decode("utf-8")
        document_xml = archive.read("word/document.xml").decode("utf-8")
    settings_root = ElementTree.fromstring(settings)
    theme_languages = settings_root.findall(qn("w:themeFontLang"))
    assert len(theme_languages) == 1
    assert theme_languages[0].get(qn("w:val")) == "vi-VN"
    assert theme_languages[0].get(qn("w:eastAsia")) == "vi-VN"
    styles_root = ElementTree.fromstring(styles)
    default_language = styles_root.find(
        f"{qn('w:docDefaults')}/{qn('w:rPrDefault')}/{qn('w:rPr')}/{qn('w:lang')}"
    )
    assert default_language is not None
    assert default_language.get(qn("w:val")) == "vi-VN"
    assert default_language.get(qn("w:eastAsia")) == "vi-VN"
    assert relationships.count("relationships/hyperlink") >= 9
    assert "https://aws.amazon.com/free/" in relationships
    assert document_xml.count("w:cantSplit") >= 78


def test_docx_guides_a_learner_from_download_to_first_lab(tmp_path: Path) -> None:
    output = tmp_path / "roadmap.docx"
    build(output)
    document = Document(output)
    text = _extract_text(document)
    for expected in (
        "Bắt đầu trong 15 phút",
        "https://github.com/quanntm1206/AWS-Cloud-Club",
        "pwsh scripts/setup.ps1 -Profile core",
        "pwsh scripts/check.ps1 -Scope bootstrap",
        ".venv\\Scripts\\python.exe scripts/run_lab.py --lab 0",
        "roadmap/weeks/week-01.md",
        "labs/lab-00-environment-and-reproducibility/README.md",
    ):
        assert expected in text


def test_docx_has_usable_navigation_and_self_assessment(tmp_path: Path) -> None:
    output = tmp_path / "roadmap.docx"
    build(output)
    document = Document(output)
    text = _extract_text(document)
    for expected in (
        "Bản đồ 6 chặng",
        "Nền tảng dữ liệu và toán",
        "AWS capstone",
        "Cách tự chấm",
        "Chưa vững",
        "Đang tiến bộ",
        "Đạt yêu cầu",
        "Vượt mong đợi",
        "Gate bắt buộc",
        "Khi bị kẹt",
        "Khi mắc kẹt",
        "Danh mục lab",
    ):
        assert expected in text
    with ZipFile(output) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")
    assert document_xml.count("w:bookmarkStart") >= 10
    assert 'w:anchor="quick-start"' in document_xml
    assert 'w:anchor="roadmap"' in document_xml


def test_docx_glossary_uses_readable_labeled_paragraphs(tmp_path: Path) -> None:
    output = tmp_path / "roadmap.docx"
    build(output)
    document = Document(output)
    glossary_heading = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text == "8. Glossary và nguồn" and paragraph.style.name == "Heading 1"
    )
    assert glossary_heading.paragraph_format.page_break_before is True
    glossary = next(table for table in document.tables if table.cell(0, 0).text == "Thuật ngữ")
    assert len(glossary.rows) == 79
    for row in glossary.rows[1:]:
        assert len(row.cells[1].paragraphs) >= 2
        assert len(row.cells[2].paragraphs) >= 3
        assert all(paragraph.text.strip() for paragraph in row.cells[1].paragraphs[:2])
        assert all(paragraph.text.strip() for paragraph in row.cells[2].paragraphs[:3])


def test_docx_adds_week_page_breaks_without_forcing_known_blank_pages(tmp_path: Path) -> None:
    output = tmp_path / "roadmap.docx"
    build(output)
    document = Document(output)
    week_headings = [
        paragraph for paragraph in document.paragraphs if re.match(r"^Tuần \d{2} - ", paragraph.text)
    ]
    assert len(week_headings) == 24
    for week, paragraph in enumerate(week_headings, start=1):
        assert paragraph.paragraph_format.page_break_before is (week not in {1, 22})
